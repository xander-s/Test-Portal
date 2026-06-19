import os
import zipfile
import shutil
import tempfile
import csv
from typing import List, Dict, Any
import openpyxl
from docx import Document
from docx.shared import Inches
from sqlalchemy.ext.asyncio import AsyncSession
import pypdf


from app.core.s3 import s3_storage
from app.models.models import Question, QuestionOption

class ImportEngineService:
    @staticmethod
    def parse_excel_or_csv(file_path: str) -> List[Dict[str, Any]]:
        """
        Parses questions from an Excel (.xlsx) or CSV template.
        """
        questions = []
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".xlsx":
            wb = openpyxl.load_workbook(file_path, read_only=True)
            sheet = wb.active
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                return []
            headers = [str(h).strip() if h else "" for h in rows[0]]
            
            for row in rows[1:]:
                if not any(row):
                    continue
                row_dict = dict(zip(headers, row))
                questions.append(row_dict)
        else:
            with open(file_path, mode="r", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    questions.append(dict(row))

        return questions

    @staticmethod
    async def process_zip_import(
        zip_file_path: str, 
        organization_id: str, 
        topic_id: str, 
        db: AsyncSession
    ) -> Dict[str, Any]:
        """
        Extracts a ZIP containing a spreadsheet and media directories.
        Uploads referenced media to MinIO and links them to the questions.
        """
        temp_dir = tempfile.mkdtemp()
        result = {"success": 0, "failed": 0, "errors": []}
        
        try:
            with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # Find the spreadsheet file (xlsx or csv)
            spreadsheet_file = None
            for root, dirs, files in os.walk(temp_dir):
                for f in files:
                    if f.endswith(('.xlsx', '.csv')) and not f.startswith('~$'):
                        spreadsheet_file = os.path.join(root, f)
                        break
            
            if not spreadsheet_file:
                raise ValueError("No valid CSV or Excel spreadsheet found inside the ZIP package.")

            raw_questions = ImportEngineService.parse_excel_or_csv(spreadsheet_file)
            
            for idx, raw_q in enumerate(raw_questions):
                try:
                    # Resolve media references relative to extraction directory
                    media_fields = ["imageUrl", "audioUrl", "videoUrl", "documentUrl"]
                    uploaded_urls = {}
                    
                    for field in media_fields:
                        val = raw_q.get(field)
                        if val:
                            local_path = os.path.join(temp_dir, str(val).strip())
                            if os.path.exists(local_path) and os.path.isfile(local_path):
                                # Determine mime type or file extension
                                filename = os.path.basename(local_path)
                                key = f"organizations/{organization_id}/questions/{filename}"
                                with open(local_path, "rb") as f:
                                    s3_url = s3_storage.upload_file(f.read(), key)
                                uploaded_urls[field] = s3_url
                            else:
                                uploaded_urls[field] = val  # Keep string if it's already an external URL
                        else:
                            uploaded_urls[field] = None

                    # Construct Question
                    question = Question(
                        topic_id=topic_id,
                        difficulty=raw_q.get("difficulty", "Medium"),
                        type=raw_q.get("type", "mcq_single"),
                        question_text=raw_q.get("text", "Empty Question"),
                        question_image=uploaded_urls.get("imageUrl"),
                        question_audio=uploaded_urls.get("audioUrl"),
                        question_video=uploaded_urls.get("videoUrl"),
                        question_document=uploaded_urls.get("documentUrl"),
                        correct_answer=raw_q.get("correctAnswer"),
                        explanation=raw_q.get("explanation"),
                        marks=float(raw_q.get("marks", 1.0)),
                        negative_marks=float(raw_q.get("negativeMarks", 0.0)),
                        organization_id=organization_id
                    )
                    db.add(question)
                    await db.flush()

                    # Add Options
                    for i in range(1, 6):
                        opt_text = raw_q.get(f"option{i}Text")
                        opt_img = raw_q.get(f"option{i}ImageUrl")
                        opt_is_correct = str(raw_q.get(f"option{i}IsCorrect", "")).lower() in ("true", "1", "yes")

                        if opt_text or opt_img:
                            # Map option images if in ZIP
                            opt_uploaded_img = None
                            if opt_img:
                                opt_img_path = os.path.join(temp_dir, str(opt_img).strip())
                                if os.path.exists(opt_img_path) and os.path.isfile(opt_img_path):
                                    filename = os.path.basename(opt_img_path)
                                    key = f"organizations/{organization_id}/options/{filename}"
                                    with open(opt_img_path, "rb") as f:
                                        opt_uploaded_img = s3_storage.upload_file(f.read(), key)
                                else:
                                    opt_uploaded_img = opt_img

                            option = QuestionOption(
                                question_id=question.id,
                                option_text=opt_text,
                                option_image=opt_uploaded_img,
                                is_correct=opt_is_correct
                            )
                            db.add(option)

                    result["success"] += 1
                except Exception as ex:
                    result["failed"] += 1
                    result["errors"].append(f"Row {idx+2}: {str(ex)}")

            await db.commit()
        finally:
            shutil.rmtree(temp_dir)

        return result

    @staticmethod
    def parse_line_options(text: str) -> List[tuple]:
        """
        Parses multiple horizontally written options from a single line.
        """
        import re
        pattern = r'(?:^|\s+)(?:\(([a-fA-F0-9])\)|([a-fA-F])[\.\)])\s+'
        matches = list(re.finditer(pattern, text))
        
        if len(matches) <= 1:
            return []
            
        options = []
        for idx, match in enumerate(matches):
            start_idx = match.end()
            end_idx = matches[idx+1].start() if idx+1 < len(matches) else len(text)
            opt_text = text[start_idx:end_idx].strip()
            opt_letter = match.group(1) or match.group(2)
            options.append((opt_letter.upper(), opt_text))
        return options

    @staticmethod
    def determine_question_type(q: Dict[str, Any]) -> str:
        """
        Infers the question type (mcq_single, mcq_multi, true_false, fill_blank).
        """
        import re
        options = q.get("options", [])
        correct_answer = q.get("correct_answer", "").strip()
        
        # 1. TrueOrFalse (true_false)
        if len(options) == 2:
            opt_texts = {o["option_text"].strip().lower() for o in options}
            if opt_texts == {"true", "false"} or opt_texts == {"yes", "no"}:
                return "true_false"
        if correct_answer.lower() in ("true", "false", "yes", "no") and len(options) <= 2:
            return "true_false"
            
        # 2. MCQMultiple (mcq_multi)
        correct_count = sum(1 for o in options if o.get("is_correct"))
        if correct_count > 1:
            return "mcq_multi"
        if len(options) > 1 and correct_answer:
            letters_found = [l.upper() for l in re.findall(r'\b([A-Za-z0-9])\b', correct_answer)]
            if len(letters_found) > 1:
                valid_letters = [chr(65 + i) for i in range(len(options))]
                if all(l in valid_letters for l in letters_found):
                    return "mcq_multi"

        # 3. FillInTheBlanks (fill_blank)
        if not options:
            if correct_answer.lower() in ("true", "false", "yes", "no"):
                return "true_false"
            return "fill_blank"
            
        # 4. MCQSingle (mcq_single)
        return "mcq_single"

    @staticmethod
    def is_roman_numeral_line(text: str) -> bool:
        """
        Checks if a line starts with a Roman numeral like I., II., III., etc.
        """
        import re
        return bool(re.match(r'^(?:[IVX]+)[\.\)]\s+(.*)', text, re.IGNORECASE))

    @staticmethod
    def parse_raw_text_to_questions(lines: List[str]) -> List[Dict[str, Any]]:
        """
        Parses questions from a list of plain text lines using regex matching for Q/Question, options, answers, and explanations.
        """
        import re
        questions = []
        current_q = None
        pending_context_lines = []
        
        for line in lines:
            text = line.strip()
            if not text:
                continue
                
            q_match = re.match(r'^(?:Q|Question)\s*[:\.\-]?\s*(.*)', text, re.IGNORECASE)
            num_match = re.match(r'^(\d+)[\.\)]\s+(.*)', text)
            
            # Check for horizontal inline options
            inline_opts = ImportEngineService.parse_line_options(text)
            if inline_opts and current_q:
                for opt_letter, opt_text in inline_opts:
                    current_q["options"].append({
                        "option_text": opt_text,
                        "is_correct": False
                    })
                continue
                
            # Restrict option labels to A-F (and parenthesized A-F/digits) to avoid matching Roman numerals
            opt_match = re.match(r'^([A-Fa-f])[\.\)]\s*(.*)', text) or re.match(r'^\(([A-Fa-f0-9])\)\s*(.*)', text)
            ans_match = re.match(r'^(?:Correct\s*Option|Correct\s*Answer|Ans|Answer|Key)\s*[:\.\-]?\s*(.*)', text, re.IGNORECASE)
            exp_match = re.match(r'^(?:Exp|Explanation)\s*[:\.\-]?\s*(.*)', text, re.IGNORECASE)
            
            if q_match:
                if current_q:
                    questions.append(current_q)
                current_q = {
                    "question_text": q_match.group(1).strip(),
                    "options": [],
                    "correct_answer": "",
                    "explanation": "",
                    "difficulty": "Medium",
                    "type": "mcq_single"
                }
                if pending_context_lines:
                    current_q["question_text"] = "\n".join(pending_context_lines) + "\n" + current_q["question_text"]
                    pending_context_lines = []
            elif num_match and not opt_match:
                if current_q:
                    questions.append(current_q)
                current_q = {
                    "question_text": num_match.group(2).strip(),
                    "options": [],
                    "correct_answer": "",
                    "explanation": "",
                    "difficulty": "Medium",
                    "type": "mcq_single"
                }
                if pending_context_lines:
                    current_q["question_text"] = "\n".join(pending_context_lines) + "\n" + current_q["question_text"]
                    pending_context_lines = []
            elif opt_match and current_q:
                opt_letter = opt_match.group(1).upper()
                opt_text = opt_match.group(2).strip()
                current_q["options"].append({
                    "option_text": opt_text,
                    "is_correct": False
                })
            elif ans_match and current_q:
                current_q["correct_answer"] = ans_match.group(1).strip()
            elif exp_match and current_q:
                current_q["explanation"] = exp_match.group(1).strip()
            elif current_q:
                # Append multi-line content
                if current_q["correct_answer"]:
                    if current_q["explanation"]:
                        current_q["explanation"] += "\n" + text
                    else:
                        pending_context_lines.append(text)
                else:
                    if not current_q["options"]:
                        current_q["question_text"] += "\n" + text
                    else:
                        current_q["options"][-1]["option_text"] += "\n" + text
            else:
                pending_context_lines.append(text)

        if current_q:
            questions.append(current_q)
            
        # Post-process correct answer match & determine question type
        for q in questions:
            # Heuristic: Extract options from the bottom of the question text lines
            # if we have no options yet but have a valid correct letter option A-F
            if not q["options"] and q["correct_answer"] and q["question_text"]:
                ans_str = q["correct_answer"].strip().upper()
                if len(ans_str) == 1 and 'A' <= ans_str <= 'F':
                    q_lines = q["question_text"].split("\n")
                    collected_opts = []
                    for line_idx in range(len(q_lines) - 1, -1, -1):
                        line_text = q_lines[line_idx].strip()
                        if not line_text:
                            continue
                        if ImportEngineService.is_roman_numeral_line(line_text) or line_text.lower().startswith("statements:"):
                            break
                        collected_opts.append(line_text)
                        if len(collected_opts) >= 6:
                            break
                    
                    required_count = ord(ans_str) - ord('A') + 1
                    if len(collected_opts) >= required_count and len(collected_opts) >= 2:
                        collected_opts.reverse()
                        q["options"] = [{"option_text": opt, "is_correct": False} for opt in collected_opts]
                        remaining_lines = q_lines[:-len(collected_opts)]
                        q["question_text"] = "\n".join(remaining_lines).strip()
            
            # Post-process matching correct options
            if q["options"] and q["correct_answer"]:
                ans_str = q["correct_answer"].strip().upper()
                letters_found = [l.upper() for l in re.findall(r'\b([A-Za-z0-9])\b', ans_str)]
                
                letters = [chr(65 + i) for i in range(len(q["options"]))]
                matched_indices = []
                for idx, opt in enumerate(q["options"]):
                    opt_letter = letters[idx]
                    if opt_letter in letters_found or opt["option_text"].upper() == ans_str:
                        matched_indices.append(idx)
                
                if matched_indices:
                    for idx in matched_indices:
                        q["options"][idx]["is_correct"] = True
                    q["correct_answer"] = ", ".join(q["options"][idx]["option_text"] for idx in matched_indices)
            
            # Auto classify type
            q["type"] = ImportEngineService.determine_question_type(q)
            
        return questions

    @staticmethod
    def parse_docx_questions(docx_file_path: str) -> List[Dict[str, Any]]:
        """
        Parses text structured questions out of a DOCX document.
        """
        doc = Document(docx_file_path)
        lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return ImportEngineService.parse_raw_text_to_questions(lines)

    @staticmethod
    def parse_pdf_questions(pdf_file_path: str) -> List[Dict[str, Any]]:
        """
        Parses text structured questions out of a PDF document.
        """
        reader = pypdf.PdfReader(pdf_file_path)
        full_text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                full_text += page_text + "\n"
        
        lines = full_text.splitlines()
        return ImportEngineService.parse_raw_text_to_questions(lines)

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extracts plain text from a Word document.
        """
        doc = Document(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extracts plain text from a PDF document.
        """
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    @staticmethod
    def parse_raw_text_with_ai(text: str) -> List[Dict[str, Any]]:
        """
        Invokes local Ollama with structured JSON prompting to parse questions.
        """
        import urllib.request
        import json
        from app.core.config import settings

        prompt = f"""You are an expert exam question parser. Analyze the following text extracted from a quiz or exam document.
Extract all questions and their choices, correct answers, explanations, difficulty, and question type.

Classify each question's type into one of the following:
- 'mcq_single': Multiple-choice question with one correct option.
- 'mcq_multi': Multiple-choice question with more than one correct option.
- 'true_false': True/False question.
- 'fill_blank': Fill in the blanks question.

Format your response strictly as a JSON object containing a "questions" list, where each item has the following structure:
{{
  "questions": [
    {{
      "question_text": "The full text of the question. For passage/context-based questions, include the passage/context text prepended to the question.",
      "type": "mcq_single" | "mcq_multi" | "true_false" | "fill_blank",
      "options": [
        {{ "option_text": "Text of the option", "is_correct": true/false }}
      ],
      "correct_answer": "For MCQ/TF questions, this must be the exact text of the correct option(s) (comma-separated if multiple). For fill-in-the-blanks, the correct phrase.",
      "explanation": "Explanation for the correct answer, if available in the text, otherwise leave empty.",
      "difficulty": "Easy" | "Medium" | "Hard" (default to "Medium"),
      "marks": 1.0,
      "negative_marks": 0.0
    }}
  ]
}}

Ensure that:
1. If the question has statements or instructions (e.g. Roman numeral lists like I, II, III), include them in the "question_text".
2. 'true_false' questions should have exactly two options: True and False.
3. 'fill_blank' questions should have empty "options" list.
4. The output must be valid JSON matching this schema exactly.

Document Text to Parse:
---
{text}
---
"""
        
        url = f"{settings.SELF_HOSTED_AI_URL}/api/chat"
        if "localhost" in url:
            url = url.replace("localhost", "127.0.0.1")

        payload = {
            "model": settings.AI_MODEL_NAME,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "stream": False,
            "format": "json"
        }

        try:
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=120) as response:
                resp_data = json.loads(response.read().decode("utf-8"))
                content = resp_data["message"]["content"]
                parsed_json = json.loads(content)
                questions = parsed_json.get("questions", [])
                
                for q in questions:
                    options = q.get("options", [])
                    if isinstance(options, list):
                        q["options"] = [
                            {
                                "option_text": str(opt.get("option_text", "")),
                                "is_correct": bool(opt.get("is_correct", False))
                            }
                            for opt in options if isinstance(opt, dict)
                        ]
                    else:
                        q["options"] = []
                        
                    q["question_text"] = str(q.get("question_text", "Empty Question")).strip()
                    q["type"] = str(q.get("type", "mcq_single")).strip().lower()
                    if q["type"] not in ("mcq_single", "mcq_multi", "true_false", "fill_blank"):
                        q["type"] = "mcq_single"
                    q["correct_answer"] = str(q.get("correct_answer", ""))
                    q["explanation"] = str(q.get("explanation", ""))
                    q["difficulty"] = str(q.get("difficulty", "Medium"))
                    try:
                        q["marks"] = float(q.get("marks", 1.0))
                    except:
                        q["marks"] = 1.0
                    try:
                        q["negative_marks"] = float(q.get("negative_marks", 0.0))
                    except:
                        q["negative_marks"] = 0.0
                
                return questions
        except Exception as e:
            raise RuntimeError(f"Ollama AI parsing failed: {str(e)}")

